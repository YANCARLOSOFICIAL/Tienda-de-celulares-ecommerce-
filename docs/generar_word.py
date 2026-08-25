"""Convierte docs/integracion-factus.md a un .docx con formato (Word).

Uso: python generar_word.py  (desde la raiz del repo, con python que tenga python-docx)
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO = Path(__file__).resolve().parent          # .../docs
MD = REPO / "integracion-factus.md"
OUT = REPO / "Integracion-Factus-Paso-a-Paso.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_BG = "F2F2F2"


def add_code(doc: Document, text: str) -> None:
    for line in text.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(18)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
    p.paragraph_format.space_after = Pt(8)


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    cols = max(len(header), *(len(r) for r in rows)) if rows else len(header)
    header = header + [""] * (cols - len(header))
    rows = [r + [""] * (cols - len(r)) for r in rows]
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", h))
        run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            text = re.sub(r"`(.+?)`", r"\1", value)
            table.rows[r].cells[c].text = text
    doc.add_paragraph()


def flush_paragraph(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(buffer).strip()
    if text:
        p = doc.add_paragraph()
        # negritas **...** y código `...`
        parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
            elif part:
                p.add_run(part)
    buffer.clear()


def convert(md_path: Path, out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)

    in_code = False
    code_lines: list[str] = []
    para_buffer: list[str] = []
    table_header: list[str] | None = None
    table_rows: list[list[str]] = []

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            flush_paragraph(doc, para_buffer)
            if not in_code:
                in_code = True
            else:
                add_code(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            continue
        if in_code:
            code_lines.append(line)
            continue

        # tablas markdown
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                continue  # separador
            if table_header is None:
                table_header = cells
            else:
                table_rows.append(cells)
            continue
        if table_header is not None:
            add_table(doc, table_header, table_rows)
            table_header, table_rows = None, []

        if not line.strip():
            flush_paragraph(doc, para_buffer)
            continue

        if line.startswith("# "):
            flush_paragraph(doc, para_buffer)
            h = doc.add_heading(line[2:].strip(), level=0)
            for run in h.runs:
                run.font.color.rgb = ACCENT
            continue
        if line.startswith("## "):
            flush_paragraph(doc, para_buffer)
            h = doc.add_heading(line[3:].strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = ACCENT
            continue
        if line.startswith("### "):
            flush_paragraph(doc, para_buffer)
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if re.match(r"^\s*[-*] \[ \]", line):
            flush_paragraph(doc, para_buffer)
            doc.add_paragraph("☐ " + line.strip()[6:], style="List Bullet")
            continue
        if re.match(r"^\s*[-*] ", line):
            flush_paragraph(doc, para_buffer)
            doc.add_paragraph(re.sub(r"^(\s*)[-*] ", "", line), style="List Bullet")
            continue
        if re.match(r"^\s*\d+\.", line):
            flush_paragraph(doc, para_buffer)
            doc.add_paragraph(re.sub(r"^\s*\d+\. ", "", line), style="List Number")
            continue
        if line.startswith(">"):
            flush_paragraph(doc, para_buffer)
            p = doc.add_paragraph()
            run = p.add_run(re.sub(r"[`*]", "", line.lstrip("> ").strip()))
            run.italic = True
            run.font.color.rgb = ACCENT
            continue
        if line.strip() in ("---", "--- "):
            flush_paragraph(doc, para_buffer)
            doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        para_buffer.append(line.strip())

    flush_paragraph(doc, para_buffer)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Generado: {out_path}")


if __name__ == "__main__":
    convert(MD, OUT)
